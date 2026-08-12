from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE

OUT = "/Users/macbook/Documents/IMSSA MEDIA EVALUATION/Weather_App_Code_Explanation.docx"
BLUE = RGBColor(21, 101, 192)
DARK = RGBColor(23, 35, 60)
MUTED = RGBColor(96, 112, 133)
LIGHT = "EAF2FB"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.72); sec.bottom_margin = Inches(.72)
sec.left_margin = Inches(.8); sec.right_margin = Inches(.8)
sec.header_distance = Inches(.3); sec.footer_distance = Inches(.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.12
for name, size, color, before, after in [
    ("Title", 26, DARK, 0, 6), ("Subtitle", 12, MUTED, 0, 18),
    ("Heading 1", 17, BLUE, 15, 7), ("Heading 2", 13, BLUE, 10, 4),
    ("Heading 3", 11, DARK, 8, 3)]:
    st = styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.color.rgb=color
    st.font.bold = name != "Subtitle"; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
    st.paragraph_format.keep_with_next=True

code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
code.font.name="Consolas"; code.font.size=Pt(8.5); code.font.color.rgb=RGBColor(32, 49, 70)
code.paragraph_format.left_indent=Inches(.22); code.paragraph_format.right_indent=Inches(.12)
code.paragraph_format.space_before=Pt(3); code.paragraph_format.space_after=Pt(7)

header=sec.header.paragraphs[0]; header.text="WEATHER APP  |  CODE REFERENCE"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.name="Calibri"; r.font.size=Pt(8); r.font.color.rgb=MUTED
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')
def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,"D9EAF7"); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c)
        for r in c.paragraphs[0].runs: r.bold=True; r.font.color.rgb=DARK; r.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cells[i])
            for p in cells[i].paragraphs: p.paragraph_format.space_after=Pt(0)
            if widths: cells[i].width=Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t
def bullet(text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.28); p.paragraph_format.first_line_indent=Inches(-.18); p.add_run(text)
def num(text):
    p=doc.add_paragraph(style="List Number"); p.paragraph_format.left_indent=Inches(.3); p.paragraph_format.first_line_indent=Inches(-.2); p.add_run(text)
def codeblock(text):
    p=doc.add_paragraph(style="Code"); p.add_run(text)
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F3F6F9'); pPr.append(shd)
def note(label,text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.15); p.paragraph_format.right_indent=Inches(.15)
    p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(8)
    r=p.add_run(label+": "); r.bold=True; r.font.color.rgb=BLUE; p.add_run(text)
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),LIGHT); pPr.append(shd)

p=doc.add_paragraph(style="Title"); p.add_run("Weather App\nCode Explanation & Viva Guide")
doc.add_paragraph("Kotlin Android application using WeatherAPI, Retrofit and Gson", style="Subtitle")
note("Purpose","A lecturer-ready explanation of the submitted application: architecture, files, request flow, JSON mapping, validation, security, testing and likely oral questions.")

doc.add_heading("1. Project Overview",1)
doc.add_paragraph("The Weather App lets a user enter a city and retrieve its current weather from WeatherAPI.com. It sends an HTTP GET request, receives JSON, converts the JSON into Kotlin objects, and displays the city, temperature, condition, humidity and wind speed dynamically.")
for x in ["Validates empty input before calling the API.","Shows a progress indicator while the request is running.","Handles invalid cities, invalid API credentials, request limits and network failures.","Keeps the API key out of GitHub by loading it from local.properties."]: bullet(x)

doc.add_heading("2. End-to-End Application Flow",1)
for x in ["User enters a city name.","Search button calls submitSearch().","Input and API-key configuration are validated.","Retrofit constructs and sends a GET request asynchronously.","WeatherAPI returns JSON.","Gson converts JSON into WeatherResponse objects.","MainActivity displays the result or an appropriate error message."]: num(x)
codeblock("User → MainActivity → WeatherApiService → Retrofit → WeatherAPI\nWeatherAPI JSON → Gson → WeatherResponse → MainActivity → UI")

doc.add_heading("3. Main Project Files",1)
table(["File","Responsibility"],[
    ("activity_main.xml","Defines the complete user interface."),("MainActivity.kt","Controls validation, requests, responses, loading and displayed data."),
    ("WeatherApiService.kt","Declares the GET endpoint and query parameters."),("RetrofitClient.kt","Builds the shared Retrofit service."),
    ("WeatherResponse.kt","Maps the JSON response to Kotlin data classes."),("AndroidManifest.xml","Adds Internet permission and declares MainActivity."),
    ("app/build.gradle.kts","Adds dependencies and exposes the local API key through BuildConfig."),("local.properties","Stores the key locally; it must not be committed."),
    ("colors.xml / drawable XML","Defines the visual palette and shaped backgrounds.")], [2.0,4.3])

doc.add_heading("4. REST API Details",1)
table(["Item","Value"],[
    ("Provider","WeatherAPI.com"),("Base URL","https://api.weatherapi.com/v1/"),("Endpoint","current.json"),
    ("HTTP method","GET"),("Required parameters","key (API key), q (city)"),("Response format","JSON")], [2.0,4.3])
codeblock("https://api.weatherapi.com/v1/current.json?key=API_KEY&q=Colombo")
doc.add_paragraph("GET is appropriate because the app reads current data and does not create, update or delete server data.")

doc.add_heading("5. Gradle Configuration",1)
codeblock('implementation("com.squareup.retrofit2:retrofit:2.11.0")\nimplementation("com.squareup.retrofit2:converter-gson:2.11.0")')
bullet("Retrofit creates and executes HTTP requests.")
bullet("converter-gson converts JSON into Kotlin data objects.")
note("Current project","Coroutines and Coil are also declared but are not used by the present callback-based implementation. They may be removed or described as planned enhancements.")

doc.add_heading("6. AndroidManifest.xml",1)
codeblock('<uses-permission android:name="android.permission.INTERNET" />')
doc.add_paragraph("This normal permission allows network communication. It is placed inside <manifest> but outside <application>, and it does not require a runtime permission dialog.")

doc.add_heading("7. JSON Models: WeatherResponse.kt",1)
codeblock('data class WeatherResponse(\n    val location: Location,\n    val current: CurrentWeather\n)')
doc.add_paragraph("The top-level JSON has location and current objects, so WeatherResponse has properties with those exact names.")
table(["JSON path","Kotlin property","Type"],[
    ("location.name","location.name","String"),("current.temp_c","temperatureCelsius","Double"),
    ("current.condition.text","condition.text","String"),("current.humidity","humidity","Int"),("current.wind_kph","windSpeedKph","Double")],[2.4,2.6,1.3])
codeblock('@SerializedName("temp_c") val temperatureCelsius: Double\n@SerializedName("wind_kph") val windSpeedKph: Double')
doc.add_paragraph("@SerializedName maps snake_case JSON names to readable camelCase Kotlin names.")

doc.add_heading("8. API Interface: WeatherApiService.kt",1)
codeblock('@GET("current.json")\nfun getCurrentWeather(\n    @Query("key") apiKey: String,\n    @Query("q") city: String\n): Call<WeatherResponse>')
bullet("@GET selects the HTTP method and relative endpoint.")
bullet("@Query adds URL query parameters and safely encodes their values.")
bullet("Call<WeatherResponse> tells Retrofit the expected parsed response type.")

doc.add_heading("9. Retrofit Client",1)
codeblock('Retrofit.Builder()\n    .baseUrl(BASE_URL)\n    .addConverterFactory(GsonConverterFactory.create())\n    .build()\n    .create(WeatherApiService::class.java)')
doc.add_paragraph("RetrofitClient is an object, so it is a singleton. The client is shared instead of rebuilt for every search. The service uses by lazy, meaning it is created only on first use. Retrofit requires the base URL to end with a slash.")

doc.add_heading("10. API-Key Handling",1)
codeblock('WEATHER_API_KEY=YOUR_PRIVATE_KEY')
doc.add_paragraph("The key is stored in local.properties. Gradle reads it with Properties and generates BuildConfig.WEATHER_API_KEY. MainActivity reads that generated constant when making a request.")
codeblock('buildConfigField("String", "WEATHER_API_KEY", "\\\"$weatherApiKey\\\"")')
note("Security limitation","This prevents accidental GitHub exposure, but BuildConfig values remain inside the compiled APK. A production app should usually call a secure backend rather than embedding a valuable secret in the mobile client.")

doc.add_heading("11. User Interface",1)
for x in ["ScrollView supports smaller screens and scrolling.","LinearLayout stacks controls vertically.","cityEditText accepts one city name and exposes a Search keyboard action.","progressBar is hidden until a request starts.","errorTextView is hidden until validation or an API error occurs.","weatherResultContainer is hidden initially and shown only after successful data arrives."]: bullet(x)
doc.add_heading("Important view IDs",2)
table(["ID","Use"],[
    ("cityEditText","City input"),("searchButton","Starts search"),("progressBar","Loading feedback"),("errorTextView","Error feedback"),
    ("weatherResultContainer","Whole result card"),("cityTextView","Returned city"),("temperatureTextView","Celsius value"),
    ("conditionTextView","Weather description"),("humidityTextView","Humidity"),("windTextView","Wind speed")],[2.4,3.9])

doc.add_heading("12. MainActivity.kt",1)
doc.add_heading("onCreate()",2)
doc.add_paragraph("onCreate runs when Android creates the screen. It loads activity_main.xml, applies system-bar insets, binds XML views with findViewById, and registers click/keyboard listeners.")
codeblock('searchButton.setOnClickListener { submitSearch() }')

doc.add_heading("submitSearch(): validation",2)
codeblock('val city = cityEditText.text.toString().trim()')
doc.add_paragraph("trim removes leading and trailing spaces. If the result is empty, the app clears old weather, shows an error and returns before any network request.")
codeblock('if (BuildConfig.WEATHER_API_KEY.isBlank()) {\n    showError("API key is not configured.")\n    return\n}')

doc.add_heading("requestWeather(): asynchronous request",2)
codeblock('RetrofitClient.weatherApi\n    .getCurrentWeather(BuildConfig.WEATHER_API_KEY, city)\n    .enqueue(object : Callback<WeatherResponse> { ... })')
doc.add_paragraph("enqueue performs the request asynchronously, so the Android main thread remains responsive. Network work must not block the UI thread.")

doc.add_heading("onResponse() and onFailure()",2)
bullet("onResponse runs whenever an HTTP response is received, including error status codes.")
bullet("response.isSuccessful checks for a 2xx response.")
bullet("response.body() returns the Gson-parsed WeatherResponse.")
bullet("onFailure runs when no usable HTTP response is obtained, such as no Internet, DNS failure or timeout.")

doc.add_heading("displayWeather()",2)
codeblock('weatherResultContainer.visibility = View.VISIBLE\ncityTextView.text = weather.location.name\ntemperatureTextView.text = String.format(Locale.getDefault(), "%.1f°C", ...)')
doc.add_paragraph("The values are dynamic and come from the API. %.1f formats temperature to one decimal place. Labels are fixed UI text, but the measurements are not hard-coded.")

doc.add_heading("Helper methods",2)
table(["Method","Role"],[
    ("clearWeather()","Hides the card and removes old values."),("showError()","Displays the supplied message."),("clearError()","Removes an old error."),
    ("setLoading()","Shows/hides ProgressBar and disables/enables Search.")],[2.2,4.1])

doc.add_heading("13. Error Handling",1)
table(["Situation","Handling"],[
    ("Empty city","Validated locally; no API request."),("HTTP 400","City not found message."),("HTTP 401","API key missing or invalid."),
    ("HTTP 403","Access denied or request limit reached."),("Other HTTP status","Displays the returned status code."),("Network failure","Connection error from onFailure().")],[2.3,4.0])
doc.add_paragraph("The app also hides stale weather when an error occurs, so users do not mistake an old result for the current search.")

doc.add_heading("14. Testing Checklist",1)
table(["Test","Expected result"],[
    ("Colombo","Live Colombo weather"),("Kandy","Live Kandy weather"),("London","Live London weather"),
    ("InvalidCity123456","City-not-found message"),("Empty input","Enter-city validation"),("Internet disabled","Network error")],[2.5,3.8])

doc.add_heading("15. Likely Viva Questions",1)
qas=[
    ("Why Retrofit?","It reduces networking boilerplate, supports annotations/query parameters, callbacks and converter integration."),
    ("Why Gson?","It converts JSON into strongly typed Kotlin data classes."),("Why GET?","The application retrieves data without modifying server state."),
    ("Why data classes?","They concisely represent structured response data and provide useful generated methods."),
    ("Why asynchronous execution?","It keeps the user interface responsive while the network request completes."),
    ("Difference between onResponse and onFailure?","onResponse handles received HTTP responses; onFailure handles communication/conversion failures with no usable response."),
    ("Why hide the result card initially?","No result exists until a successful request completes."),
    ("Are values hard-coded?","Only labels are fixed; weather values come from the JSON response."),
    ("Is the API key fully secure?","It is excluded from GitHub, but the APK still contains it; a production backend is safer."),
    ("What does @SerializedName do?","It maps JSON names such as temp_c to Kotlin properties such as temperatureCelsius.")]
for q,a in qas:
    p=doc.add_paragraph(); r=p.add_run(q+" "); r.bold=True; r.font.color.rgb=BLUE; p.add_run(a)

doc.add_heading("16. Short Presentation Script",1)
doc.add_paragraph("Our application is a Kotlin Android Weather App integrated with WeatherAPI.com. The user enters a city, and the app validates the input before Retrofit sends an asynchronous GET request. WeatherAPI returns JSON, which Gson maps into Kotlin data classes. The app dynamically displays city, temperature, condition, humidity and wind speed. It also handles empty input, invalid cities, invalid credentials, request limits and network failures. The API key is stored in local.properties so it is not committed to GitHub.")

doc.add_heading("17. Final Submission Notes",1)
for x in ["Never submit or screenshot the real API key.","Confirm local.properties is ignored by Git.","Capture Postman GET + 200 OK + JSON without exposing credentials.","Capture successful app output and at least one error case.","Be ready to explain the full request path and each model field."]: bullet(x)

doc.core_properties.title="Weather App Code Explanation and Viva Guide"
doc.core_properties.subject="Android REST API integration using Kotlin, Retrofit, Gson and WeatherAPI"
doc.core_properties.author="Weather App Project Team"
doc.save(OUT)
print(OUT)
