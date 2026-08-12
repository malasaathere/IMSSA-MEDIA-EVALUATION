from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('/Users/macbook/Documents/IMSSA MEDIA EVALUATION/outputs/IMSSA_UI_UX_Feature_Specification.docx')
SCREENSHOT = Path('/var/folders/9z/mwmvblnj24x28sbym6kc_wbm0000gn/T/codex-clipboard-84140138-8fd9-4b37-96be-33a5b10a6428.png')
VIDEO_SCREENSHOT = Path('/var/folders/9z/mwmvblnj24x28sbym6kc_wbm0000gn/T/codex-clipboard-3087b172-efa1-45dd-891c-e2675afa021c.png')
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = '0B2545'; BLUE = '2E74B5'; CYAN = '0E7490'; TEAL = '047857'
LIGHT = 'E8EEF5'; PALE = 'F4F7FA'; GOLD = 'B7791F'; RED = '9B1C1C'; GRAY = '5B6573'; WHITE = 'FFFFFF'

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35); section.footer_distance = Inches(0.35)

def font(run, size=10.5, bold=False, color=NAVY, italic=False):
    run.font.name = 'Aptos'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Aptos')
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Aptos')
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Aptos'; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.15
for name, size, color, before, after in [
    ('Heading 1', 17, BLUE, 16, 8), ('Heading 2', 13.5, BLUE, 11, 5), ('Heading 3', 11.5, NAVY, 8, 3)]:
    st = styles[name]; st.font.name = 'Aptos Display'; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color); st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
for name in ['List Bullet', 'List Number']:
    st = styles[name]; st.font.name = 'Aptos'; st.font.size = Pt(10.5); st.font.color.rgb = RGBColor.from_string(NAVY)
    st.paragraph_format.left_indent = Inches(0.35); st.paragraph_format.first_line_indent = Inches(-0.18)
    st.paragraph_format.space_after = Pt(3); st.paragraph_format.line_spacing = 1.12
if 'UI Label' not in styles:
    st = styles.add_style('UI Label', WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = 'Aptos'; st.font.size = Pt(9); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(CYAN)
    st.paragraph_format.space_before = Pt(5); st.paragraph_format.space_after = Pt(2)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr(); tcMar = tc.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tc.append(tcMar)
    for tag, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        el = tcMar.find(qn('w:'+tag))
        if el is None: el = OxmlElement('w:'+tag); tcMar.append(el)
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa')

def set_cell_text(cell, text, bold=False, color=NAVY, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.08
    font(p.add_run(str(text)), size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)

def table(headers, rows, widths=None, header_fill=BLUE, font_size=8.8):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    t.style = 'Table Grid'
    if widths is None: widths = [6.7/len(headers)] * len(headers)
    for i, h in enumerate(headers):
        t.columns[i].width = Inches(widths[i]); set_cell_text(t.rows[0].cells[i], h, True, WHITE, 9.0); shade(t.rows[0].cells[i], header_fill)
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, NAVY, font_size)
            if r_i % 2: shade(cells[i], PALE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

def para(text='', bold_lead=None, italic=False, color=NAVY, size=10.5, after=5, align=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.15
    if align is not None: p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), size, True, color)
        font(p.add_run(text[len(bold_lead):]), size, False, color, italic)
    else: font(p.add_run(text), size, False, color, italic)
    return p

def bullets(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet'); font(p.add_run(item), 10.2)

def steps(items):
    for item in items:
        p = doc.add_paragraph(style='List Number'); font(p.add_run(item), 10.2)

def callout(title, body, fill=LIGHT, accent=BLUE):
    t = doc.add_table(rows=1, cols=1); t.autofit=False; t.columns[0].width=Inches(6.65)
    c=t.cell(0,0); shade(c, fill); margins(c, 130, 160, 130, 160); c.text=''
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3); font(p.add_run(title), 10.5, True, accent)
    p=c.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.12; font(p.add_run(body), 9.7, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def page_break(): doc.add_page_break()

def heading(text, level=1): doc.add_heading(text, level=level)

# Running furniture
hp = section.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run('IMSSA MEDIA EVALUATION  |  UI/UX SPECIFICATION'), 8.5, True, GRAY)
fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run('Internal product design reference  •  August 2026  •  '), 8.2, False, GRAY)
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); fp._p.append(fld)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(12)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
font(p.add_run('PRODUCT DESIGN SPECIFICATION'), 10, True, CYAN)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(9)
font(p.add_run('IMSSA Media Evaluation Platform'), 27, True, NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
font(p.add_run('Complete UI feature, content, interaction and UX requirements'), 14, False, BLUE)
if not SCREENSHOT.exists():
    raise FileNotFoundError(f'Cover screenshot not found: {SCREENSHOT}')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
run=p.add_run(); inline_shape=run.add_picture(str(SCREENSHOT), width=Inches(6.45))
doc_pr = inline_shape._inline.docPr
doc_pr.set('name', 'IMSSA Media workspace login interface')
doc_pr.set('descr', 'Desktop login screen showing the IMSSA Media brand panel and four role-based workspace options.')
callout('Purpose', 'A build-ready reference for a role-based media-production workspace covering planning, assignment, creation, review, approval, communication, reporting and administration.', PALE, TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4)
font(p.add_run('Version 1.1  •  Desktop, tablet and mobile web  •  August 2026'), 8.8, True, GRAY)
para('Design principle: show the next meaningful action clearly, preserve context across roles, and never make status depend on color alone.', italic=True, color=GRAY, size=9.3, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

page_break()
heading('1. Product vision and UX outcomes')
para('The platform should replace fragmented spreadsheets, chats and file handovers with one traceable workflow. The interface must help every user answer four questions immediately: What needs attention? Who owns it? When is it due? What happens next?')
heading('1.1 Success outcomes',2)
table(['Outcome','UX requirement','How to measure'],[
 ['Fast orientation','A role-specific home screen with urgent work, due dates and primary action above the fold.','User identifies next task in under 10 seconds.'],
 ['Clear ownership','Every task and plan item shows assignee, creator, event and status.','No unlabelled or ambiguous ownership.'],
 ['Safe handovers','Uploads, revisions, review decisions and comments remain attached to the task.','Users do not need external chat to reconstruct history.'],
 ['Low training burden','Consistent labels, controls, status language and action placement.','New users complete core flows without assistance.'],
 ['Mobile viability','Core view, update, upload, review and message actions work at 360 px.','No clipped controls or horizontal page scrolling.'],
 ['Trustworthy feedback','Every save, error, permission block and background sync has a visible state.','No silent failures.']
],[1.25,3.35,2.1])
heading('1.2 Experience principles',2)
bullets([
 'Role first: each dashboard prioritizes the job that role performs most often.',
 'Progressive disclosure: cards show essential status; dialogs and detail pages reveal full context.',
 'Action proximity: place decisions beside the content they affect.',
 'Recognition over recall: show event, owner, status definitions and previous actions instead of relying on memory.',
 'Recoverability: confirm destructive actions, preserve drafts and provide retry paths.',
 'Evidence-based status: “approved,” “posted” and “completed” must link to the supporting version or activity.',
])

page_break()
heading('2. Users, roles and permissions')
table(['Role','Primary goal','Default landing','Key actions'],[
 ['Marketing Coordinator','Plan campaigns and create/assign work.','Marketing Workspace','Create tasks; manage plans; import sheet data; view team capacity.'],
 ['Designer / Video Editor','Complete assigned creative work.','My Work','Acknowledge task; upload draft; submit revision; message; use checklists.'],
 ['Media Director','Review work and issue decisions.','Review Inbox','Inspect version; annotate; approve; request revision; provide feedback.'],
 ['Chief Coordinator','Monitor delivery and workload.','Analytics','Review KPIs, bottlenecks, overdue work and team capacity.'],
 ['Administrator','Manage access and platform configuration.','Administration','Approve requests; add/deactivate users; manage roles; inspect system health.'],
 ['Multi-role user','Switch between authorized work contexts.','Last-used role or role chooser','Change role context without signing out.']
],[1.35,1.7,1.2,2.45],font_size=8.5)
callout('Permission rule', 'Hide actions users cannot perform, but show a clear Access Denied explanation when they follow a saved or shared link. Never expose restricted data briefly during loading.')

heading('3. Information architecture and global shell')
heading('3.1 Primary navigation',2)
table(['Navigation item','Audience','Contents / behavior'],[
 ['Home / Marketing','Marketing coordinators','Priority summary, task board, capacity and recent activity.'],
 ['Marketing Plan','Marketing coordinators','Campaign tabs, filters, plan table/cards, edit and import actions.'],
 ['My Work','Designers and editors','Assigned task queue, active task, upload and revision workflow.'],
 ['Review Inbox','Media directors','Submitted versions requiring review and decision.'],
 ['Analytics','Chief coordinators and admins','Delivery KPIs, capacity, status and event performance.'],
 ['Calendar','All authenticated roles','Deadlines and publishing dates, filterable by event/person/status.'],
 ['Administration','Admins','Users, requests, roles, configuration and health.'],
 ['Profile','All users','Identity, roles, preferences, security and session controls.']
],[1.35,1.55,3.8])
heading('3.2 Persistent shell requirements',2)
bullets([
 'Desktop: collapsible left sidebar, top bar, page title, optional contextual actions and content area.',
 'Mobile: compact top bar with menu button; navigation opens as a labelled drawer and returns focus when closed.',
 'Global search: search tasks and plan items by title, event, assignee and status; group results by content type.',
 'Notification center: unread count, filter by type, mark read and deep-link to the affected task/version.',
 'Team chat: floating launcher on desktop; bottom-sheet or full-screen view on mobile; unread badge and timestamps.',
 'Role switcher: appear only for multi-role users and preserve their last context.',
 'System status: show sync degradation or service outage without blocking unrelated work.',
])
heading('3.3 Global feedback states',2)
table(['State','Required UI'],[
 ['Loading','Use skeletons matching final layout; preserve page title and navigation.'],
 ['Empty','Explain why it is empty and offer the next valid action, e.g. “No tasks assigned yet.”'],
 ['Error','Plain-language message, retry action, reference code and safe path back.'],
 ['Offline / degraded','Persistent non-modal banner; queue safe drafts where supported.'],
 ['Saved','Short confirmation near the action; avoid disruptive success dialogs.'],
 ['Unsaved changes','Visible dirty state and leave-page confirmation.'],
 ['Permission denied','State required role and offer navigation to an allowed workspace.']
],[1.55,5.15],font_size=9.2)

heading('4. Authentication, registration and profile')
heading('4.1 Login',2)
table(['UI area','Required details'],[
 ['Role portal cards','Marketing Coordinator, Chief Coordinator, Designers & Editors, Media Directors; show icon and one-line purpose.'],
 ['Credential input','Label “4-digit passkey”; numeric keypad on mobile; masked value; show/hide control; no placeholder-only label.'],
 ['Primary action','“Sign in”; disabled only while empty or submitting; loading text “Signing in…”.'],
 ['Errors','Invalid passkey, inactive account, network failure and unavailable role; do not reveal whether a passkey belongs to another person.'],
 ['Support','Link to request access and contact administrator.']
],[1.5,5.2])
heading('4.2 Registration / access request',2)
table(['Field','Validation and UX'],[
 ['Full name','Required; trim whitespace; 2–100 characters.'],
 ['Email','Required; format validation; explain it is used for account communication.'],
 ['Requested role','Use canonical role names: Designer, Video Editor, Marketing Coordinator, Media Director, Chief Coordinator.'],
 ['Reason for access','Required short explanation; character counter; 20–500 characters.'],
 ['Submission state','Confirmation with request status and expected next step; prevent duplicate submissions.']
],[1.5,5.2])
heading('4.3 Profile and preferences',2)
bullets(['Name, avatar, email, status, timezone and all assigned roles.', 'Notification preferences: in-app alerts, email alerts and deadline reminders.', 'Security area: change/reset passkey, active session and sign-out action.', 'Accessibility preference: reduced motion and optional high-contrast mode.', 'Do not display passkeys or secrets in profile, admin tables, exports or logs.'])

heading('5. Marketing Coordinator UI')
heading('5.1 Marketing Workspace dashboard',2)
table(['Section','Content','Primary interactions'],[
 ['Priority strip','Overdue, due today, waiting assignment, in review.','Select metric to filter task board.'],
 ['Task board','To Do, In Progress, In Review, Done; task count per column.','Open task; create task; filter by event/assignee/priority.'],
 ['Capacity panel','Eligible designers/editors, active task count, capacity indicator.','Select a person to filter or prefill assignee.'],
 ['Recent activity','Task creation, uploads, reviews, approvals and status changes.','Open referenced item; filter by event.'],
 ['Quick actions','Create task, import sheet, open plan, open calendar.','Actions remain visible on desktop; overflow menu on mobile.']
],[1.35,3.15,2.2],font_size=8.7)
heading('5.2 Create Task dialog',2)
table(['Field / control','Content requirement','UX behavior'],[
 ['Title','Clear deliverable name; example: “Opening Ceremony Banner”.','Required, 3–255 characters; live validation.'],
 ['Description / brief','Purpose, audience, format, mandatory content and references.','Required recommendation; multiline; character count.'],
 ['Event','HackX 11.0, HackX Jr 9.0 or Exposition 2026.','Required; searchable selection.'],
 ['Work type','Poster, social post, video, animation, photo, copy or other configured type.','Required; influences upload guidance.'],
 ['Priority','Low, Medium, High.','Default Medium; always pair color with text.'],
 ['Assignee','Eligible designer or video editor with capacity.','Required; show active workload and capacity warning.'],
 ['Deadline','Date and time in Asia/Colombo.','Required; prevent past dates; show relative time.'],
 ['References','Links or files supporting the brief.','Optional; show allowed formats and upload progress.'],
 ['Submit','Create and assign.','Review summary, prevent duplicate submission, show success link.']
],[1.35,3.1,2.25],font_size=8.3)
heading('5.3 Marketing Plan screen',2)
bullets([
 'Campaign selector for HackX 11.0, HackX Jr 9.0 and Exposition 2026, with record counts.',
 'Search and filters for type, designer, design status, caption status, final status, platform and share date.',
 'Desktop table columns: post title/description, type, designer, writer, design status, caption status, final status, handover date, due/share dates and platform.',
 'Mobile card layout: title and urgent status first; secondary details in expandable sections.',
 'Row actions: view, edit, convert to task, generate caption draft, open calendar item and view sync state.',
 'Bulk/import: assign or update selected rows, export/sync, map CSV headers, resolve duplicates and preview normalized dates before confirmation.',
])

page_break()
heading('6. Designer and Video Editor UI')
if not VIDEO_SCREENSHOT.exists():
    raise FileNotFoundError(f'Video editing screenshot not found: {VIDEO_SCREENSHOT}')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3)
run=p.add_run(); inline_shape=run.add_picture(str(VIDEO_SCREENSHOT), width=Inches(6.45))
doc_pr = inline_shape._inline.docPr
doc_pr.set('name', 'Exposition event video frame')
doc_pr.set('descr', 'Warm-toned Exposition event footage showing an attendee holding the published magazine, with the Exposition issue mark in the upper-right corner.')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(9)
font(p.add_run('Video-editing reference frame: branded event footage, safe logo placement and publication-focused storytelling.'), 8.8, False, GRAY, True)
heading('6.1 My Work landing',2)
table(['Area','Required content'],[
 ['Work summary','Assigned, due soon, in progress, revision requested and waiting review counts.'],
 ['Task queue','Title, event, priority, deadline, current status and next required action.'],
 ['Active task','Full brief, references, workflow stepper, assignee and creator.'],
 ['Capacity','Personal active count and limit; explain why new work may be blocked.'],
 ['Checklist','Designer pack items relevant to quality and brand compliance.']
],[1.45,5.25])
heading('6.2 Task lifecycle interactions',2)
steps([
 'Open an assigned task and acknowledge ownership.',
 'Review the brief, required dimensions, references and deadline.',
 'Start work; status becomes In Progress and the change is logged.',
 'Upload a draft with a short version note and file preview.',
 'Run the revision/quality checklist before submission.',
 'Submit for review; lock the submitted version while allowing a new revision to be created later.',
 'If revision is requested, show director feedback beside the affected version and start the next revision round.',
 'When approved, show the approved version, approval timestamp and any performance feedback.'
])
heading('6.3 Upload Draft panel',2)
table(['Element','Required behavior'],[
 ['Drop zone / file picker','Keyboard accessible; supported formats and maximum size shown before selection.'],
 ['Upload queue','Filename, type, size, progress, cancel/retry and failure reason.'],
 ['Preview','Image/video preview where supported; fallback icon and safe download link.'],
 ['Version note','Required concise description of changes.'],
 ['Submit action','Disabled until upload completes and checklist requirements are met.'],
 ['Confirmation','Show version number, review status and link back to task.']
],[1.55,5.15])

heading('7. Media Director Review UI')
heading('7.1 Review Inbox',2)
table(['Area','Details'],[
 ['Queue list','Thumbnail, task title, event, creator, submitted time, deadline and revision round.'],
 ['Filters','Event, work type, assignee, priority, age and revision round.'],
 ['Sort','Oldest waiting, nearest deadline, highest priority or newest submission.'],
 ['Preview canvas','Fit, zoom, pan, full screen and file metadata; preserve task context.'],
 ['Context panel','Brief, previous versions, checklist result, discussion and status history.'],
 ['Decision bar','Approve or Request Revision; sticky on desktop and mobile.']
],[1.45,5.25])
heading('7.2 Review and annotation behavior',2)
bullets([
 'Annotations must be anchored to a coordinate or timestamp and show author, time, status and replies.',
 'Use numbered markers on the canvas and a synchronized comment list; selecting either highlights both.',
 'Allow resolve/reopen; preserve resolved comments in history.',
 'Autosave comment drafts and show saving/saved/error state.',
 'For video, support timestamped comments and frame capture when technically available.',
])
heading('7.3 Approval decision',2)
table(['Decision','Required inputs','Result'],[
 ['Approve','Optional positive note; confirm selected version.','Version becomes Approved; task progresses to Completed/Approved; notify relevant users.'],
 ['Request Revision','Required actionable feedback; optional annotated points; severity/priority.','Version becomes Revision Requested; increment revision workflow; notify assignee.'],
 ['Performance Feedback','Strengths, improvements and general suggestions.','Saved separately from production decision and visible to authorized roles.']
],[1.25,2.8,2.65],font_size=8.7)
callout('Decision safety', 'The approval action must name the exact version and task. Do not use a generic “Yes/No” confirmation. Revision feedback should be specific enough for the designer to act without requesting clarification.', 'FDECEC', RED)

heading('8. Chief Coordinator Analytics UI')
heading('8.1 Executive overview',2)
table(['Metric','Definition','Interaction'],[
 ['Active work','Tasks not completed or cancelled.','Open filtered task list.'],
 ['Completed / approved','Tasks with approved evidence in the selected period.','Drill into event/person details.'],
 ['In review','Submitted versions awaiting director action.','Sort by wait time.'],
 ['In progress','Acknowledged or actively worked tasks.','Filter by event and assignee.'],
 ['Overdue','Open tasks whose deadlines have passed.','Show owner and escalation path.'],
 ['On-time rate','Completed before deadline divided by completed tasks.','Show calculation definition and sample size.']
],[1.25,3.5,1.95],font_size=8.7)
heading('8.2 Visualizations',2)
bullets([
 'Designer workload: horizontal bars with active count, capacity limit and overdue segment.',
 'Status distribution: labelled stacked bar preferred over a color-only pie.',
 'Event delivery trend: tasks created, completed and overdue by week.',
 'Review turnaround: median time from submission to decision, with outliers visible.',
 'Plan readiness: design, caption and final-status completion by campaign.',
 'Every chart requires a text summary, legend, tooltip, filter state and downloadable data table.',
])

heading('9. Administrator UI')
table(['Screen / module','Required features'],[
 ['User management','Search, role/status filters, user detail, role assignment, activate/deactivate and audit context.'],
 ['Add user','Name, email, one or more roles, timezone and generated onboarding/reset flow. Never reveal stored passkeys.'],
 ['Access requests','Request details, duplicate indicator, approve/reject with reason and confirmation.'],
 ['Role management','Canonical role definitions and permissions; protect removal when assigned users exist.'],
 ['System health','Appwrite connection, function status, Sheet/Calendar sync, queue failures and last successful backup.'],
 ['Audit log','Actor, action, target, timestamp and outcome; filters and export for authorized admins.'],
 ['Data tools','Safe import preview, export, duplicate detection and non-destructive migration status.']
],[1.55,5.15])

heading('10. Calendar, search, notifications and communication')
heading('10.1 Calendar',2)
bullets([
 'Month, week, agenda and mobile list views.',
 'Combine task deadlines and marketing publishing dates with distinct icons and labels.',
 'Filters: event, assignee, content type, status and source; provide “Clear all”.',
 'Event detail shows source record, owner, date, status and direct navigation.',
 'Conflict indicators for overloaded assignees or multiple high-priority deadlines.',
 'Visible sync status: synced, pending, failed or not connected; manual retry when allowed.',
 'Use Asia/Colombo consistently while preserving explicit timezone metadata.',
])
heading('10.2 Global search',2)
table(['Result type','Preview details','Actions'],[
 ['Task','Title, event, assignee, status and deadline.','Open task; copy link.'],
 ['Marketing plan item','Campaign, post title, platform, owner and share date.','Open/edit item; convert to task.'],
 ['User','Name, roles and status.','Open profile when permitted.'],
 ['File/version','Task, version number, status and uploader.','Open review/version view.']
],[1.35,3.7,1.65])
heading('10.3 Notifications',2)
bullets(['Group by Today, Earlier and Read.', 'Types include assignment, deadline reminder, draft submission, revision request, approval, mention and sync failure.', 'Each notification names the actor, object and time and deep-links to the exact context.', 'Provide mark read, mark all read and preference link; avoid relying on browser notifications alone.'])
heading('10.4 Task discussion and team chat',2)
bullets(['Task discussion is task-scoped and should be used for decisions and clarifications.', 'Global team chat is general communication and must not replace task history.', 'Show sender name/avatar, timestamp, delivery state and unread marker.', 'Support keyboard submission, multiline text, links and safe attachment references.', 'Provide empty, loading, failed-send and retry states.'])

heading('11. Design system')
table(['Token / component','Specification'],[
 ['Color','Navy for structure, blue/teal for actions and status accents, neutral surfaces; meet WCAG contrast.'],
 ['Typography','Aptos/Inter/system sans; body 16 px desktop and mobile; minimum supporting text 14 px.'],
 ['Spacing','4 px base grid; common gaps 8, 12, 16, 24 and 32 px.'],
 ['Touch targets','Minimum 44×44 CSS px; 48×48 preferred for primary mobile actions.'],
 ['Buttons','Primary, secondary, tertiary and destructive; include hover, focus, disabled and loading states.'],
 ['Forms','Persistent labels, help text, inline errors and summary on submit when multiple fields fail.'],
 ['Status chips','Text + icon + color; canonical labels shared across screens.'],
 ['Cards','Use for compact records; avoid nested cards and excessive elevation.'],
 ['Tables','Sticky header, responsive columns, row actions, empty state and accessible header associations.'],
 ['Dialogs','Use only for focused tasks; full-screen on narrow mobile for complex forms.']
],[1.5,5.2],font_size=9.0)
heading('11.1 Canonical task statuses',2)
table(['Status','Meaning','Recommended treatment'],[
 ['Draft','Not yet assigned.','Neutral gray, document icon.'],['Assigned / To Do','Owner assigned; work not started.','Blue, assignment icon.'],['In Progress','Work actively underway.','Cyan, progress icon.'],['Ready for Review / In Review','Submitted and awaiting decision.','Purple, eye icon.'],['Revision Requested','Changes required.','Amber, revision icon.'],['Approved / Completed','Accepted final version.','Green, check icon.'],['Cancelled','No further work.','Gray/red outline, cancelled icon.']
],[1.65,2.75,2.3],font_size=8.7)

page_break()
heading('12. Responsive UX requirements')
table(['Breakpoint','Layout expectations'],[
 ['Mobile: 320–639 px','Single column; drawer navigation; full-width forms; bottom action bar; card alternatives to wide tables.'],
 ['Tablet: 640–1023 px','Two-column where useful; collapsible sidebar; review panels stack or resize.'],
 ['Desktop: 1024–1439 px','Persistent sidebar; multi-panel workspaces; sticky filters/actions.'],
 ['Large: 1440 px+','Constrain reading width; use extra space for context panels, not stretched text.']
],[1.55,5.15])
heading('12.1 Mobile acceptance rules',2)
bullets([
 'No page-level horizontal scrolling at 320 px.',
 'Primary action is reachable with one hand and remains visible during long forms/reviews.',
 'Tables transform to cards or controlled horizontal regions with clear affordance.',
 'Dialogs become full-screen when content exceeds a simple confirmation.',
 'File upload supports camera/gallery/file chooser as available.',
 'Keyboard opening must not cover the active input or submit action.',
 'Charts provide a readable summary and data list; tooltips cannot be hover-only.',
])

heading('13. Accessibility and inclusive design')
bullets([
 'Meet WCAG 2.2 AA for contrast, keyboard access, focus visibility, labels and error identification.',
 'Use one H1 per page and logical heading order; landmarks for header, navigation, main and complementary regions.',
 'All controls require accessible names; icon-only controls require tooltips and programmatic labels.',
 'Do not encode status or priority using color alone.',
 'Maintain visible focus with at least a 2 px contrasting outline.',
 'Announce asynchronous success/error updates through an ARIA live region without stealing focus.',
 'Provide captions/transcripts or text alternatives for review media where applicable.',
 'Respect reduced-motion settings; avoid flashing and unnecessary animation.',
 'Allow browser zoom to 200% without loss of content or function.',
])

page_break()
heading('14. Validation, privacy and security UX')
table(['Area','Requirement'],[
 ['Passkeys','Mask input; rate-limit attempts; never display or export stored passkeys; provide admin reset flow.'],
 ['Role enforcement','Validate on server and client; UI hiding is not authorization.'],
 ['Uploads','Validate type/size, scan where available, prevent executable content and use expiring access links.'],
 ['Destructive actions','Confirm target and consequence; prefer deactivate/archive over permanent deletion.'],
 ['Auditability','Record actor, action, target, timestamp and result for privileged changes.'],
 ['Personal data','Show only data needed for the user’s role; avoid email/passkey exposure in broad team views.'],
 ['Errors','Do not expose API keys, stack traces, database IDs or internal configuration to users.']
],[1.45,5.25])

heading('15. Screen-level content checklist')
table(['Every record/detail screen should show','Every interactive screen should include'],[
 ['Clear title and content type\nEvent/campaign context\nOwner and creator\nCanonical status\nDeadline/date and timezone\nLast updated / relevant history\nEvidence or linked version',
  'Primary action\nSecondary/back action\nLoading state\nEmpty state\nValidation state\nPermission state\nError + retry state\nSave/success feedback\nMobile behavior'],
], [3.35,3.35], font_size=9.4)

heading('16. MVP, next phase and future scope')
table(['Priority','Include'],[
 ['MVP / required','Role-based shell; login/request access; marketing task board; task creation; My Work; upload/version submission; Review Inbox; approval/revision; plans; calendar; notifications; profile; admin users; complete loading/error/empty/mobile states.'],
 ['Phase 2','Persistent annotations/replies; robust task/global chat; Google Sheet and Calendar sync dashboard; audit log UI; advanced search; data exports; richer analytics.'],
 ['Future','AI pre-check with human override; automated caption suggestions; video timestamp review; configurable workflows; advanced workload forecasting; offline draft support.']
],[1.35,5.35])
callout('Scope rule', 'A feature should not be shown as complete unless its data is persisted, permission-checked, recoverable after refresh and represented in activity/history where appropriate.', PALE, TEAL)

heading('17. End-to-end acceptance scenarios')
steps([
 'A marketing coordinator creates a high-priority task for Exposition, sees an assignee capacity warning, assigns it and receives confirmation.',
 'The designer sees the task immediately, opens the brief on mobile, acknowledges it and begins work.',
 'The designer uploads a valid draft, completes the checklist and submits version 1 for review.',
 'The media director receives a notification, opens the exact version, leaves anchored feedback and requests revision.',
 'The designer sees the revision reason in context, uploads version 2 and resubmits.',
 'The director approves version 2; the approved evidence and timestamps appear in the task history.',
 'The chief coordinator sees updated completion and review-turnaround metrics with a drill-down to the task.',
 'An administrator can verify the relevant audit activity without viewing passkeys or unrelated sensitive data.'
])

heading('18. Designer handoff checklist')
bullets([
 'Create desktop and mobile frames for every primary route and critical dialog.',
 'Provide component variants for hover, focus, pressed, disabled, loading, error and success.',
 'Document design tokens, spacing, breakpoints, status colors/icons and chart palette.',
 'Include realistic data density using the three campaigns and role structure.',
 'Prototype the create-task, upload-review-revision and access-request flows.',
 'Annotate permissions and conditional visibility in the design file.',
 'Run keyboard, contrast, 200% zoom and 360 px layout checks before development handoff.',
 'Confirm all field names and status labels match the database/API contract.',
])

heading('Appendix A. Route inventory')
table(['Route','Screen','Authorized roles'],[
 ['/login','Sign in','Public'],['/register','Request access','Public'],['/','Marketing Workspace','Marketing Coordinator'],['/marketing-plan','Marketing Plan','Marketing Coordinator'],['/designer','My Work','Designer, Video Editor'],['/director','Review Inbox','Media Director'],['/analytics','Analytics','Chief Coordinator, Admin'],['/calendar','Calendar','All authenticated roles'],['/admin','Administration','Admin'],['/admin/users','User Requests / Users','Admin'],['/profile','Profile','All authenticated roles']
],[1.35,2.9,2.45],font_size=8.8)

heading('Appendix B. Data-backed UI entities')
table(['Entity','UI representation'],[
 ['Users / roles','Identity, role switcher, assignee controls, capacity and permissions.'],['Events','Campaign selector, filters, colors and event context.'],['Marketing plan items','Plan table/cards, publishing calendar and progress.'],['Tasks / assignments','Boards, queues, owner and deadlines.'],['Status history','Timeline and analytics evidence.'],['Versions / files','Upload queue, previews and approved evidence.'],['Reviews / annotations','Decision history and anchored feedback.'],['Messages / notifications','Task discussion, team chat and attention center.'],['Preferences','Notification and accessibility controls.']
],[1.75,4.95])

doc.core_properties.title = 'IMSSA Media Evaluation Platform UI/UX Feature Specification'
doc.core_properties.subject = 'Complete interface, interaction, responsive and accessibility requirements'
doc.core_properties.author = 'IMSSA Media Platform Team'
doc.core_properties.keywords = 'IMSSA, UI, UX, product design, media workflow'
doc.save(OUT)
print(OUT)
